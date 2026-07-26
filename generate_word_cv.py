import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Colors
DARK_BLUE = RGBColor(20, 78, 196) # #144ec4
TEXT_DARK = RGBColor(15, 23, 42)  # #0f172a
TEXT_MUTED = RGBColor(71, 85, 105) # #475569

doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Styles
styles = doc.styles

# Normal style (Calibri, 11pt)
style_normal = styles['Normal']
font_normal = style_normal.font
font_normal.name = 'Calibri'
font_normal.size = Pt(11)
font_normal.color.rgb = TEXT_DARK

# Define custom styles safely
def get_or_create_style(name, style_type, base_style_name=None):
    if name in styles:
        return styles[name]
    style = styles.add_style(name, style_type)
    if base_style_name:
        style.base_style = styles[base_style_name]
    return style

# Title style
style_title = get_or_create_style('CV_Title', WD_STYLE_TYPE.PARAGRAPH)
font_title = style_title.font
font_title.name = 'Calibri'
font_title.size = Pt(24)
font_title.bold = True
font_title.color.rgb = DARK_BLUE
style_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_title.paragraph_format.space_after = Pt(4)

# Subtitle style
style_sub = get_or_create_style('CV_Subtitle', WD_STYLE_TYPE.PARAGRAPH)
font_sub = style_sub.font
font_sub.name = 'Calibri'
font_sub.size = Pt(14)
font_sub.color.rgb = TEXT_MUTED
style_sub.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_sub.paragraph_format.space_after = Pt(20)

# Section Header style
style_header = get_or_create_style('CV_Header', WD_STYLE_TYPE.PARAGRAPH)
font_header = style_header.font
font_header.name = 'Calibri'
font_header.size = Pt(14)
font_header.bold = True
font_header.color.rgb = DARK_BLUE
style_header.paragraph_format.space_before = Pt(16)
style_header.paragraph_format.space_after = Pt(6)

# Experience Title style
style_exp_title = get_or_create_style('CV_Exp_Title', WD_STYLE_TYPE.PARAGRAPH)
font_exp = style_exp_title.font
font_exp.name = 'Calibri'
font_exp.size = Pt(12)
font_exp.bold = True
font_exp.color.rgb = TEXT_DARK
style_exp_title.paragraph_format.space_after = Pt(2)
style_exp_title.paragraph_format.space_before = Pt(8)

# Date/Company style
style_exp_sub = get_or_create_style('CV_Exp_Sub', WD_STYLE_TYPE.PARAGRAPH)
font_exp_sub = style_exp_sub.font
font_exp_sub.name = 'Calibri'
font_exp_sub.size = Pt(10)
font_exp_sub.italic = True
font_exp_sub.color.rgb = DARK_BLUE
style_exp_sub.paragraph_format.space_after = Pt(4)

# Create Document Content

# Header
doc.add_paragraph("PAPE ALIOUNE SENE", style='CV_Title')
doc.add_paragraph("Développeur Full-Stack & UX/UI Designer", style='CV_Subtitle')

# Contact info
contact_text = "📱 +221 77 XXX XX XX  |  ✉️ senepapealioune0@gmail.com\n🌍 Dakar, Sénégal  |  🔗 github.com/Alioune205  |  🔗 linkedin.com/in/papealiounesene"
p_contact = doc.add_paragraph(contact_text)
p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_contact.runs[0].font.size = Pt(10)
p_contact.runs[0].font.color.rgb = TEXT_MUTED
p_contact.paragraph_format.space_after = Pt(20)

# Profile
doc.add_paragraph("PROFIL PROFESSIONNEL", style='CV_Header')
profile_text = (
    "Étudiant en BTS Informatique et fondateur de PAS Digital Studio, je suis passionné par le développement "
    "d'applications web, mobiles et IA à fort impact. Je crée des solutions innovantes pour l'écosystème "
    "africain : registres civils intelligents, e-santé, civic-tech et souveraineté numérique. Mon objectif : "
    "participer activement au développement socio-économique du Sénégal par la technologie."
)
doc.add_paragraph(profile_text)

# Experience
doc.add_paragraph("PROJETS & RÉALISATIONS", style='CV_Header')

def add_experience(title, date_role, desc, tags):
    doc.add_paragraph("💻 " + title, style='CV_Exp_Title')
    doc.add_paragraph(date_role, style='CV_Exp_Sub')
    p = doc.add_paragraph(desc)
    p.paragraph_format.space_after = Pt(4)
    p_tags = doc.add_paragraph("⚙️ Technologies : " + tags)
    p_tags.runs[0].font.size = Pt(9.5)
    p_tags.runs[0].font.color.rgb = TEXT_MUTED
    p_tags.paragraph_format.space_after = Pt(12)

add_experience(
    "CFP Abéné — Site Institutionnel", 
    "2024 — Centre de Formation Professionnelle d'Abéné",
    "Conception et développement du site web officiel : 10 filières de formation, actualités, formulaire d'inscription en ligne, partenaires internationaux et intégration cartographique.",
    "WordPress, Elementor, SEO"
)

add_experience(
    "Teranga Civil / SUNU CIVIL",
    "2024 — GovTech (Registre civil numérique)",
    "Plateforme GovTech centralisant les données d'état civil : gestion des rôles, sécurité avancée, architecture full-stack scalable et tableau de bord analytique.",
    "Django, React, PostgreSQL"
)

add_experience(
    "DiabetoConseil",
    "2024 — Mobile Health",
    "Application mobile intégrant un assistant IA pour le suivi glycémique, les conseils personnalisés et la prise de rendez-vous médical.",
    "Flutter, Groq IA (Llama 3.3), Django"
)

add_experience(
    "WaxalëPay",
    "2024 — Civic Tech",
    "Portail permettant aux citoyens de vérifier leur éligibilité aux programmes sociaux de l'État (PNBSF, CMU) avec un chatbot IA d'accompagnement.",
    "UI/UX, Chatbot IA, Figma"
)

# Skills
doc.add_paragraph("COMPÉTENCES & OUTILS", style='CV_Header')
skills_text = (
    "• Backend : Django REST Framework, PostgreSQL, Python\n"
    "• Frontend : React, JavaScript, Angular\n"
    "• Mobile : Flutter\n"
    "• IA & Data : Groq IA, Machine Learning, Power BI\n"
    "• Design : Figma, UI/UX, Photoshop, Montage vidéo\n"
    "• Autres : Git/GitHub, n8n, WordPress, Odoo"
)
p_skills = doc.add_paragraph(skills_text)
p_skills.paragraph_format.line_spacing = 1.3

# Education
doc.add_paragraph("FORMATION", style='CV_Header')

def add_education(degree, school_year):
    doc.add_paragraph("🎓 " + degree, style='CV_Exp_Title')
    doc.add_paragraph(school_year, style='CV_Exp_Sub')

add_education(
    "BTS Informatique — Développement d'Applications",
    "Établissement Supérieur, Dakar, Sénégal (2023 — En cours)"
)

add_education(
    "Baccalauréat — Série Scientifique (S)",
    "Lycée, Sénégal (2023)"
)

# Languages & Interests
doc.add_paragraph("LANGUES & INTÉRÊTS", style='CV_Header')
lang_int_text = (
    "🗣️ Langues : Français (Courant), Anglais (Intermédiaire), Wolof (Natif)\n"
    "🌍 Intérêts : Civic-Tech & Impact Social, Intelligence Artificielle, Design, Football"
)
doc.add_paragraph(lang_int_text)

# Save
doc.save("C:/Users/senep/Desktop/PORLIO/Pape_Alioune_Sene_3.docx")
print("DOCX created successfully.")
